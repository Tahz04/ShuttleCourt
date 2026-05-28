import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from copy import deepcopy

def set_font(run, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(docx.oxml.ns.qn('w:rFonts'))
    if rFonts is None:
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(docx.oxml.ns.qn('w:cs'), 'Times New Roman')

def create_table(doc, headers, rows_data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Set headers
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font(run, bold=True)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
    # Add data
    for row_data in rows_data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run)
    return table

def add_heading_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(text)
    set_font(run, bold=True)
    return p

def add_empty_para(doc):
    return doc.add_paragraph()

def main():
    doc = docx.Document('Bao_Cao_ShuttleCourt.docx')
    
    # Find the target paragraph
    target_p = None
    for p in doc.paragraphs:
        if 'Kịch bản kiểm thử (Test Cases)' in p.text:
            target_p = p
            break
            
    if not target_p:
        print("ERROR: Cannot find target paragraph")
        return
        
    # Table 1: Unit Test
    ut_headers = ["Mã TC", "Tên kịch bản / Chức năng", "Đầu vào (Input)", "Kết quả kỳ vọng", "Thực tế"]
    ut_data = [
        ["UT-01", "Tính khoảng cách Haversine", "Tọa độ GPS hợp lệ (Lat, Lng của 2 điểm)", "Trả về khoảng cách chính xác (km)", "Pass"],
        ["UT-02", "Xác thực định dạng Email", "Chuỗi \"test@gmail.com\"", "Trả về True (Hợp lệ)", "Pass"],
        ["UT-03", "Mã hóa mật khẩu (Bcrypt)", "Chuỗi \"password123\"", "Trả về chuỗi Hash 60 ký tự", "Pass"],
        ["UT-04", "Tạo JWT Token", "User ID, Role = 'user'", "Trả về chuỗi Token có thể giải mã hợp lệ", "Pass"],
        ["UT-05", "Validate thời gian đặt sân", "Thời gian kết thúc < Thời gian bắt đầu", "Báo lỗi \"Thời gian không hợp lệ\"", "Pass"]
    ]
    p_ut = add_heading_para(doc, "Bảng 3.1. Kịch bản Kiểm thử Đơn vị (Unit Test)")
    tbl_ut = create_table(doc, ut_headers, ut_data)
    emp_ut = add_empty_para(doc)
    
    # Table 2: Feature Test
    ft_headers = ["Mã TC", "Tên kịch bản / Chức năng", "Đầu vào (Input)", "Kết quả kỳ vọng", "Thực tế"]
    ft_data = [
        ["FT-01", "Đăng nhập sai mật khẩu", "Email đúng, sai password", "Hiển thị thông báo \"Sai thông tin đăng nhập\"", "Pass"],
        ["FT-02", "Tìm kiếm sân theo tên", "Nhập từ khóa \"Sân A\"", "Hiển thị danh sách các sân có tên chứa \"Sân A\"", "Pass"],
        ["FT-03", "Chủ sân thêm sân mới", "Nhập đầy đủ thông tin sân hợp lệ", "Lưu CSDL thành công, sân xuất hiện trên bản đồ", "Pass"],
        ["FT-04", "Gửi yêu cầu đặt sân", "Chọn khung giờ trống hợp lệ", "Tạo đơn 'Chờ duyệt', gửi thông báo cho Owner", "Pass"],
        ["FT-05", "Duyệt yêu cầu đặt sân", "Owner click nút 'Duyệt'", "Đơn chuyển trạng thái 'Đã duyệt', báo cho User", "Pass"]
    ]
    p_ft = add_heading_para(doc, "Bảng 3.2. Kịch bản Kiểm thử Chức năng (Feature Test)")
    tbl_ft = create_table(doc, ft_headers, ft_data)
    emp_ft = add_empty_para(doc)
    
    # Table 3: Concurrency Test
    ct_headers = ["Mã TC", "Tên kịch bản / Chức năng", "Đầu vào (Input)", "Kết quả kỳ vọng", "Thực tế"]
    ct_data = [
        ["CT-01", "Hai người đặt cùng 1 sân, giờ", "User A & B gọi API đặt sân đồng thời", "Một người thành công, người kia báo lỗi \"Sân đã được đặt\"", "Pass"],
        ["CT-02", "Hai người join kèo (còn 1 slot)", "User A & B gọi API join kèo đồng thời", "Một người vào kèo, người kia báo lỗi \"Kèo đã đủ người\"", "Pass"],
        ["CT-03", "Admin khóa TK khi User thao tác", "Admin ban User, User sau đó gọi API", "API trả về lỗi 401 Unauthorized, từ chối thao tác", "Pass"]
    ]
    p_ct = add_heading_para(doc, "Bảng 3.3. Kịch bản Kiểm thử Tương tranh (Concurrency Test)")
    tbl_ct = create_table(doc, ct_headers, ct_data)
    emp_ct = add_empty_para(doc)
    
    # Move elements to target
    # Order: p_ut, tbl_ut, emp_ut, p_ft, tbl_ft, emp_ft, p_ct, tbl_ct, emp_ct
    elements_to_move = [
        p_ut._element, tbl_ut._element, emp_ut._element,
        p_ft._element, tbl_ft._element, emp_ft._element,
        p_ct._element, tbl_ct._element, emp_ct._element
    ]
    
    current_anchor = target_p._element
    for el in elements_to_move:
        current_anchor.addnext(el)
        current_anchor = el
        
    doc.save('Bao_Cao_ShuttleCourt.docx')
    print("Successfully added 3 tables for Unit, Feature, and Concurrency Tests")

if __name__ == '__main__':
    main()
